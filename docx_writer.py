import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from image_handler import save_images_for_docx
from math_handler import math_to_docx_paragraph


def html_to_docx(md_text, output_path, md_file_dir=None):
    if md_file_dir is None:
        md_file_dir = os.getcwd()

    doc = Document()
    _setup_styles(doc)

    lines = md_text.split('\n')
    i = 0
    images_cache_dir = os.path.join(md_file_dir, '.md_converter_images')
    image_list = []

    while i < len(lines):
        line = lines[i]

        if line.startswith('```'):
            lang = line[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            _add_code_block(doc, '\n'.join(code_lines), lang)
            i += 1
            continue

        if line.startswith('|') and '|' in line[1:]:
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            _add_table(doc, table_lines)
            continue

        heading_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            _add_heading(doc, text, level)
            i += 1
            continue

        if line.startswith('> '):
            quote_lines = []
            while i < len(lines) and lines[i].startswith('> '):
                quote_lines.append(lines[i][2:])
                i += 1
            _add_blockquote(doc, '\n'.join(quote_lines))
            continue

        if re.match(r'^[-*_]{3,}\s*$', line):
            i += 1
            continue

        if line.startswith('---'):
            i += 1
            continue

        ul_match = re.match(r'^(\s*)([-*+])\s+(.*)', line)
        if ul_match:
            indent = len(ul_match.group(1))
            text = ul_match.group(3).strip()
            _add_list_item(doc, text, 'bullet', indent)
            i += 1
            continue

        ol_match = re.match(r'^(\s*)\d+\.\s+(.*)', line)
        if ol_match:
            indent = len(ol_match.group(1))
            text = ol_match.group(2).strip()
            _add_list_item(doc, text, 'number', indent)
            i += 1
            continue

        if line.strip() == '':
            i += 1
            continue

        _add_paragraph(doc, line)
        i += 1

    doc.save(output_path)


def _setup_styles(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    for level in range(1, 7):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Microsoft YaHei'
        heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def _add_heading(doc, text, level):
    level = min(level, 4)
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def _add_paragraph(doc, text):
    para = doc.add_paragraph()
    _add_inline_runs(para, text)


def _add_inline_runs(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*|__.*?__|(?<!\*)\*(?!\*)(.*?)(?<!\*)\*(?!\*)|~~.*?~~|`[^`]+`|!\[.*?\]\(.*?\)|\[.*?\]\(.*?\))', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('__') and part.endswith('__'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('~~') and part.endswith('~~'):
            run = paragraph.add_run(part[2:-2])
            run.font.strike = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            shading = run.element.get_or_add_rPr()
            shading_elm = shading.makeelement(qn('w:shd'), {
                qn('w:val'): 'clear',
                qn('w:color'): 'auto',
                qn('w:fill'): 'F0F0F0',
            })
            shading.append(shading_elm)
        elif part.startswith('!['):
            img_match = re.match(r'!\[(.*?)\]\((.*?)\)', part)
            if img_match:
                alt = img_match.group(1)
                src = img_match.group(2)
                if not src.startswith('http'):
                    full = os.path.normpath(src)
                    if os.path.isfile(full):
                        paragraph.add_run().add_picture(full, width=Inches(5))
                    else:
                        paragraph.add_run(f'[Image: {alt}]')
                else:
                    paragraph.add_run(f'[Image: {alt}]')
        elif part.startswith('[') and '](' in part:
            link_match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if link_match:
                text_content = link_match.group(1)
                url = link_match.group(2)
                run = paragraph.add_run(text_content)
                run.font.color.rgb = RGBColor(0x03, 0x66, 0xD6)
                run.underline = True
        else:
            math_to_docx_paragraph(None, paragraph, part)


def _add_code_block(doc, code, lang):
    if lang:
        p = doc.add_paragraph()
        run = p.add_run(f'Language: {lang}')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    pPr = p._element.get_or_add_pPr()
    shading = pPr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F5F5F5',
    })
    pPr.append(shading)

    pf = p.paragraph_format
    pf.left_indent = Cm(0.5)
    pf.right_indent = Cm(0.5)

    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)


def _add_blockquote(doc, text):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    left_bdr = pBdr.makeelement(qn('w:left'), {
        qn('w:val'): 'single',
        qn('w:sz'): '12',
        qn('w:space'): '4',
        qn('w:color'): 'CCCCCC',
    })
    pBdr.append(left_bdr)
    pPr.append(pBdr)

    pf = p.paragraph_format
    pf.left_indent = Cm(1)

    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True


def _add_table(doc, table_lines):
    parsed = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip('|').split('|')]
        parsed.append(cells)

    if len(parsed) < 2:
        return

    if all(re.match(r'^[-:]+$', c) for c in parsed[1]):
        data = [parsed[0]] + parsed[2:]
    else:
        data = parsed

    if not data:
        return

    rows = len(data)
    cols = max(len(r) for r in data)

    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            if j < cols:
                cell = table.cell(i, j)
                cell.text = cell_text
                if i == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True


def _add_list_item(doc, text, list_type, indent):
    if list_type == 'bullet':
        p = doc.add_paragraph(style='List Bullet')
    else:
        p = doc.add_paragraph(style='List Number')

    if indent > 0:
        pf = p.paragraph_format
        pf.left_indent = Cm(1.27 * (indent // 2 + 1))

    _add_inline_runs(p, text)
